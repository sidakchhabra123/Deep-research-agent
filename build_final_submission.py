from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).parent
DOCX_PATH = OUT_DIR / "FINAL_SUBMISSION.docx"
MD_PATH = OUT_DIR / "FINAL_SUBMISSION.md"


MARKDOWN = """# Deep Research Agent: Final Submission

## Index
1. Design Note (Architecture & Strategy)
2. Evaluation Methodology & Harness Results
3. Setup Instructions, Links, & Examples

---

## 1. Design Note

### Target Users & Problem Solved

The Deep Research Agent is built for researchers, engineers, students, and technical decision-makers who need answers that are current, inspectable, and grounded in source material. Standard large language models are useful for synthesis, but they can hallucinate facts, rely on stale training data, and present fluent answers without showing where claims came from. That is a poor fit for research workflows where the user must verify evidence, compare sources, and understand uncertainty.

This project solves that problem by turning a normal question into a deterministic research workflow. The agent searches the live web, fetches pages, extracts useful text, prunes noisy material, and only then asks the LLM to synthesize an answer. Every factual claim is expected to use bracketed citations such as [1] and [2], which map directly to retrieved sources in the UI. The result is not just an answer, but a transparent research artifact: search queries, opened URLs, selected snippets, citations, and final response are all persisted in SQLite session history.

### Definition of "Deep Research"

In this implementation, "Deep Research" means a deterministic multi-step orchestration pipeline rather than a single prompt. The agent does not simply ask an LLM to answer from memory. It executes a structured process:

1. Plan targeted search queries from the user question.
2. Search the web using Tavily.
3. Fetch and extract readable HTML content using a pure-Python fetch layer with trafilatura-style extraction.
4. Score, prune, and diversify evidence under a token budget using tiktoken.
5. Synthesize a concise cited answer from only the selected web context.
6. Persist the full session, turn metadata, and source details in SQLite.

This design keeps the Planner, Search, Fetch, and Context Builder nodes in English to preserve retrieval quality and reasoning consistency. Multilingual support is added with a Translation Sandwich: user input in Hindi or Tamil is translated into English with Sarvam, the research pipeline runs in English, and the final answer is translated back to the selected language.

### Success Metrics

The agent is optimized against four core quality metrics:

- **Citation Integrity Rate:** Whether the final answer uses strict bracketed citations, such as [1] and [2], for factual claims. This measures whether answers are grounded rather than free-floating LLM assertions.
- **Source Diversity Score:** The number of unique domains represented in the citation map. This helps avoid over-reliance on a single source and improves robustness for comparison or conflict questions.
- **Uncertainty Handling:** Whether the agent refuses or qualifies answers when the retrieved evidence is insufficient. The expected behavior is to explicitly state that sufficient evidence could not be found rather than inventing an answer.
- **Conflict Resolution Accuracy:** Whether the agent detects and explains disagreement across sources. For conflicting evidence, the agent must state "Sources disagree:" and objectively summarize both sides.

### Data Flow and Components

The system is built completely from scratch in pure Python, without LangChain or similar orchestration frameworks. The key components are:

- **Planner Node:** Converts the user question and optional session history into up to three targeted search queries. The planner is explicitly instructed to decompose multi-hop, comparison, and multi-entity questions into distinct sub-queries rather than repeating keyword variants.
- **Search/Fetch Layer:** Uses Tavily for web search, then fetches pages and extracts readable content. This layer records fetch success, failure, domains, titles, word counts, and opened URLs.
- **Context Builder:** Scores pages by relevance, search score, recency, keyword density, and domain diversity. It aggressively prunes low-value snippets, penalizes duplicate domains, and ensures the citation map includes only snippets that survived final pruning and were actually placed in the context window.
- **Orchestrator:** Coordinates the full generator-based pipeline. It yields structured streaming steps to Streamlit, including planning, search, fetch results, context readiness, retry events, conflict warnings, answer tokens, and completion metadata.
- **Streamlit UI:** Renders a commercial-style Agent Workflow log, research plan expander, session history, detailed session modal, sources, and turn details.
- **SQLite Session Store:** Persists sessions, message history, turn metadata, search queries, URLs opened, snippets selected, final answers, and context breakdowns.
- **Translation Sandwich:** Uses the Sarvam API to support English, Hindi, and Tamil. Inputs are translated to English before planning/search, and completed English answers are translated back after generation.

### Risks, Limitations, & Future Improvements

The agent improves reliability but still has practical constraints. Some sites return 403 errors or block scraping, which can reduce source diversity. Very large or broad questions can exhaust context budgets even with pruning. Web search results may contain SEO spam, duplicated content, or pages with weak extraction quality. External APIs also introduce rate limits and latency, especially when multiple searches, LLM calls, and translation calls are involved.

Two important future improvements are:

1. **Iterative ReAct Retry Loop:** Expand the current retry mechanism into a multi-iteration ReAct loop that can inspect gaps, refine queries, fetch new evidence, and continue until confidence thresholds are met or a clear refusal is warranted.
2. **Local Vector Database:** Replace purely lexical snippet scoring with a local embedding index or vector database. This would improve semantic retrieval, reduce noise, and make repeated research over saved sessions faster and more accurate.

---

## 2. Evaluation Methodology & Harness Results

### The Methodology

The project includes a custom `eval/eval_harness.py` script that runs the full agent pipeline against an eight-question dataset in `eval/dataset.json`. The dataset is intentionally small but adversarial. It targets the most important edge cases for a deep research agent:

- factual lookup,
- multi-hop reasoning,
- comparison across entities,
- insufficient evidence,
- conflicting sources,
- multi-turn memory,
- strict citation formatting.

The harness instantiates the actual project components: `SessionStore`, `LLMClient`, `Planner`, `SearchClient`, `PageFetcher`, `ContextBuilder`, and `Orchestrator`. It consumes the orchestrator generator step by step, accumulating streamed `token` chunks into the final answer and extracting the `done` metadata, including `citation_map` and `pages_used`. Multi-turn evaluation reuses the same `session_id` so that a follow-up such as "that movie" can be resolved through session memory.

### The Metrics Explained

The harness computes metrics programmatically for every test case:

- **Citation Integrity:** Checks whether the final answer contains bracketed citation markers such as `[1]` or `[2]`. For successful insufficient-evidence refusals, the harness does not penalize missing citations because the correct behavior is not to cite a fabricated factual answer.
- **Source Diversity:** Counts unique source domains in the citation map. This helps measure whether the agent gathered evidence from multiple independent places.
- **Uncertainty Handling:** For insufficient-evidence questions, checks whether the answer includes refusal language such as "cannot", "insufficient", or "no information".
- **Conflict Flagging:** For conflict questions, checks whether the orchestrator yielded a conflict step or whether the answer contains "disagree" or "conflict".

### Harness Result Snapshot

| Evaluation Type | What It Tests | Expected Behavior |
|---|---|---|
| Factual | Exact known answer with sources | Provide a concise answer with bracketed citations. |
| Multi-hop | Entity linking plus prior role lookup | Break question into sub-queries and synthesize across sources. |
| Comparison | Two-model context window comparison | Cite each side and compare directly. |
| Insufficient Evidence | Fictional Atlantis GDP | Refuse with clear insufficient-evidence language. |
| Conflict | Coffee and blood pressure | Surface disagreement or nuance across sources. |
| Multi-turn | Inception director then budget | Reuse the same session context for follow-up resolution. |

### Findings Summary

The agent successfully demonstrates multi-turn memory and multi-hop reasoning by preserving session context and using planned sub-queries to resolve follow-up questions. The evaluation harness also exposed useful failure modes: citation formatting can degrade when prompts are too loose, insufficient-evidence cases require explicit refusal language, and conflict detection depends on retrieving enough independent domains.

Final hardening focused on three areas. First, the system prompt was tightened so factual sentences must end with bracketed citations. Second, the Context Builder was upgraded with domain-diversity penalties and a rescue path to avoid one-domain context. Third, the evaluation harness was corrected so a valid refusal is not unfairly penalized for missing citations. These changes align the agent with the intended research behavior: answer when evidence is sufficient, refuse when it is not, and call out conflicts when sources disagree.

---

## 3. Setup Instructions, Links, & Examples

### Links

- GitHub Repository: [Insert GitHub Link]
- Video Demo: [Insert Video Demo Link]

### Setup & Run Instructions

1. Clone the repository:

```bash
git clone [Insert GitHub Link]
cd deep_research_agent
```

2. Create and activate a Python environment:

```bash
python -m venv .venv
.venv\\Scripts\\activate
```

3. Install dependencies:

```bash
pip install -r requirements.txt
```

4. Configure the `.env` file:

```env
TAVILY_API_KEY=your_tavily_key_here
GEMINI_API_KEY=your_gemini_key_here
GROQ_API_KEY=your_groq_key_here
SARVAM_API_KEY=your_sarvam_key_here
```

5. Run the Streamlit app:

```bash
streamlit run streamlit_app.py
```

If the project is reorganized under an `app/` directory, use:

```bash
streamlit run app/streamlit_app.py
```

6. Run the evaluation harness manually:

```bash
python eval/eval_harness.py
```

The harness writes full results to:

```text
eval/results.json
```

### Example Conversations

**Example 1: Factual query with citations**

User: "What organization maintains the Unicode Standard?"

Agent: "The Unicode Standard is maintained by the Unicode Consortium [1]. The consortium coordinates standardization work for text encoding across software systems [1]."

**Example 2: Impossible query with insufficient-evidence refusal**

User: "What is the exact GDP of the fictional underwater city of Atlantis in 2024?"

Agent: "I cannot find sufficient evidence to answer this. Atlantis is fictional, and the retrieved sources do not provide a reliable real-world GDP figure. A refined search could look for fictional portrayals of Atlantis economics or worldbuilding estimates."
"""


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360, col_widths=None):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(width_dxa))
    table.autofit = False
    if col_widths:
        grid = tbl.tblGrid
        for child in list(grid):
            grid.remove(child)
        for width in col_widths:
            grid_col = OxmlElement("w:gridCol")
            grid_col.set(qn("w:w"), str(width))
            grid.append(grid_col)
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                cell.width = Pt(col_widths[idx] / 20)
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_w = tc_pr.tcW
                tc_w.type = "dxa"
                tc_w.w = col_widths[idx]
                set_cell_margins(cell)


def apply_styles(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Deep Research Agent Final Submission")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)


def add_para(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        if isinstance(item, tuple):
            label, body = item
            r = p.add_run(label)
            r.bold = True
            p.add_run(body)
        else:
            p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(32, 32, 32)


def add_metrics_table(doc):
    rows = [
        ("Factual", "Exact known answer with sources", "Provide a concise answer with bracketed citations."),
        ("Multi-hop", "Entity linking plus prior role lookup", "Break question into sub-queries and synthesize across sources."),
        ("Comparison", "Two-model context window comparison", "Cite each side and compare directly."),
        ("Insufficient Evidence", "Fictional Atlantis GDP", "Refuse with clear insufficient-evidence language."),
        ("Conflict", "Coffee and blood pressure", "Surface disagreement or nuance across sources."),
        ("Multi-turn", "Inception director then budget", "Reuse the same session context for follow-up resolution."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_width(table, 9360, [2200, 3200, 3960])
    hdr = table.rows[0].cells
    for idx, text in enumerate(["Evaluation Type", "What It Tests", "Expected Behavior"]):
        hdr[idx].text = text
        set_cell_shading(hdr[idx], "F2F4F7")
        hdr[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in hdr[idx].paragraphs:
            for run in p.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[idx])
    doc.add_paragraph()


def build_docx():
    doc = Document()
    apply_styles(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("Deep Research Agent: Final Submission")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    r = subtitle.add_run("Pure-Python research orchestration with Streamlit, SQLite, Tavily, tiktoken pruning, strict citations, evaluation harness, and Sarvam multilingual support.")
    r.italic = True
    r.font.color.rgb = RGBColor(90, 90, 90)

    doc.add_heading("Index", level=2)
    add_numbered(doc, [
        "Design Note (Architecture & Strategy)",
        "Evaluation Methodology & Harness Results",
        "Setup Instructions, Links, & Examples",
    ])

    doc.add_heading("1. Design Note", level=1)
    doc.add_heading("Target Users & Problem Solved", level=2)
    add_para(doc, "The Deep Research Agent is built for researchers, engineers, students, and technical decision-makers who need answers that are current, inspectable, and grounded in source material. Standard large language models are useful for synthesis, but they can hallucinate facts, rely on stale training data, and present fluent answers without showing where claims came from. That is a poor fit for research workflows where the user must verify evidence, compare sources, and understand uncertainty.")
    add_para(doc, "This project solves that problem by turning a normal question into a deterministic research workflow. The agent searches the live web, fetches pages, extracts useful text, prunes noisy material, and only then asks the LLM to synthesize an answer. Every factual claim is expected to use bracketed citations such as [1] and [2], which map directly to retrieved sources in the UI. The result is not just an answer, but a transparent research artifact: search queries, opened URLs, selected snippets, citations, and final response are all persisted in SQLite session history.")

    doc.add_heading('Definition of "Deep Research"', level=2)
    add_para(doc, '"Deep Research" means a deterministic multi-step orchestration pipeline rather than a single prompt. The agent does not simply ask an LLM to answer from memory. It executes a structured process:')
    add_numbered(doc, [
        "Plan targeted search queries from the user question.",
        "Search the web using Tavily.",
        "Fetch and extract readable HTML content using a pure-Python fetch layer with trafilatura-style extraction.",
        "Score, prune, and diversify evidence under a token budget using tiktoken.",
        "Synthesize a concise cited answer from only the selected web context.",
        "Persist the full session, turn metadata, and source details in SQLite.",
    ])
    add_para(doc, "This design keeps the Planner, Search, Fetch, and Context Builder nodes in English to preserve retrieval quality and reasoning consistency. Multilingual support is added with a Translation Sandwich: user input in Hindi or Tamil is translated into English with Sarvam, the research pipeline runs in English, and the final answer is translated back to the selected language.")

    doc.add_heading("Success Metrics", level=2)
    add_bullets(doc, [
        ("Citation Integrity Rate: ", "Whether the final answer uses strict bracketed citations, such as [1] and [2], for factual claims."),
        ("Source Diversity Score: ", "The number of unique domains represented in the citation map, reducing over-reliance on a single source."),
        ("Uncertainty Handling: ", "Whether the agent refuses or qualifies answers when retrieved evidence is insufficient."),
        ("Conflict Resolution Accuracy: ", "Whether the agent detects and explains disagreement across sources with objective summaries of both sides."),
    ])

    doc.add_heading("Data Flow and Components", level=2)
    add_para(doc, "The system is built completely from scratch in pure Python, without LangChain or similar orchestration frameworks. The key components are:")
    add_bullets(doc, [
        ("Planner Node: ", "Converts the user question and optional session history into up to three targeted search queries, with explicit decomposition for multi-hop and comparison tasks."),
        ("Search/Fetch Layer: ", "Uses Tavily for web search, then fetches pages and extracts readable content while recording fetch success, failure, domains, titles, word counts, and opened URLs."),
        ("Context Builder: ", "Scores pages by relevance, search score, recency, keyword density, and domain diversity, then prunes low-value snippets under a token budget."),
        ("Orchestrator: ", "Coordinates the generator-based pipeline and yields structured streaming steps to the Streamlit UI."),
        ("Streamlit UI: ", "Renders the Agent Workflow log, research plan expander, session history, detailed session modal, sources, and turn details."),
        ("SQLite Session Store: ", "Persists sessions, message history, turn metadata, search queries, URLs opened, snippets selected, final answers, and context breakdowns."),
        ("Translation Sandwich: ", "Uses Sarvam API support for English, Hindi, and Tamil while keeping core retrieval and reasoning in English."),
    ])

    doc.add_heading("Risks, Limitations, & Future Improvements", level=2)
    add_para(doc, "The agent improves reliability but still has practical constraints. Some sites return 403 errors or block scraping, which can reduce source diversity. Very large or broad questions can exhaust context budgets even with pruning. Web search results may contain SEO spam, duplicated content, or pages with weak extraction quality. External APIs also introduce rate limits and latency, especially when multiple searches, LLM calls, and translation calls are involved.")
    add_numbered(doc, [
        "Iterative ReAct Retry Loop: Expand the current retry mechanism into a multi-iteration loop that can inspect gaps, refine queries, fetch new evidence, and continue until confidence thresholds are met or a clear refusal is warranted.",
        "Local Vector Database: Replace purely lexical snippet scoring with a local embedding index or vector database for better semantic retrieval and faster repeated research over saved sessions.",
    ])

    doc.add_heading("2. Evaluation Methodology & Harness Results", level=1)
    doc.add_heading("The Methodology", level=2)
    add_para(doc, "The project includes a custom eval/eval_harness.py script that runs the full agent pipeline against an eight-question dataset in eval/dataset.json. The dataset is intentionally small but adversarial. It targets factual lookup, multi-hop reasoning, comparison across entities, insufficient evidence, conflicting sources, multi-turn memory, and strict citation formatting.")
    add_para(doc, "The harness instantiates the actual project components: SessionStore, LLMClient, Planner, SearchClient, PageFetcher, ContextBuilder, and Orchestrator. It consumes the orchestrator generator step by step, accumulating streamed token chunks into the final answer and extracting done metadata, including citation_map and pages_used. Multi-turn evaluation reuses the same session_id so that a follow-up such as 'that movie' can be resolved through session memory.")

    doc.add_heading("The Metrics Explained", level=2)
    add_bullets(doc, [
        ("Citation Integrity: ", "Checks whether the final answer contains bracketed citation markers such as [1] or [2]. Valid insufficient-evidence refusals are not penalized for missing fabricated citations."),
        ("Source Diversity: ", "Counts unique source domains in the citation map."),
        ("Uncertainty Handling: ", "Checks refusal language such as 'cannot', 'insufficient', or 'no information' for insufficient-evidence questions."),
        ("Conflict Flagging: ", "Checks whether the orchestrator yielded a conflict step or the answer contains 'disagree' or 'conflict'."),
    ])

    doc.add_heading("Harness Result Snapshot", level=2)
    add_metrics_table(doc)

    doc.add_heading("Findings Summary", level=2)
    add_para(doc, "The agent successfully demonstrates multi-turn memory and multi-hop reasoning by preserving session context and using planned sub-queries to resolve follow-up questions. The evaluation harness also exposed useful failure modes: citation formatting can degrade when prompts are too loose, insufficient-evidence cases require explicit refusal language, and conflict detection depends on retrieving enough independent domains.")
    add_para(doc, "Final hardening focused on three areas. First, the system prompt was tightened so factual sentences must end with bracketed citations. Second, the Context Builder was upgraded with domain-diversity penalties and a rescue path to avoid one-domain context. Third, the evaluation harness was corrected so a valid refusal is not unfairly penalized for missing citations. These changes align the agent with the intended research behavior: answer when evidence is sufficient, refuse when it is not, and call out conflicts when sources disagree.")

    doc.add_heading("3. Setup Instructions, Links, & Examples", level=1)
    doc.add_heading("Links", level=2)
    add_bullets(doc, [
        "GitHub Repository: [Insert GitHub Link]",
        "Video Demo: [Insert Video Demo Link]",
    ])

    doc.add_heading("Setup & Run Instructions", level=2)
    add_numbered(doc, [
        "Clone the repository.",
        "Create and activate a Python environment.",
        "Install dependencies.",
        "Configure the .env file.",
        "Run the Streamlit app.",
        "Run the evaluation harness manually.",
    ])
    add_code_block(doc, [
        "git clone [Insert GitHub Link]",
        "cd deep_research_agent",
        "python -m venv .venv",
        ".venv\\Scripts\\activate",
        "pip install -r requirements.txt",
        "",
        "TAVILY_API_KEY=your_tavily_key_here",
        "GEMINI_API_KEY=your_gemini_key_here",
        "GROQ_API_KEY=your_groq_key_here",
        "SARVAM_API_KEY=your_sarvam_key_here",
        "",
        "streamlit run streamlit_app.py",
        "# If reorganized under app/:",
        "streamlit run app/streamlit_app.py",
        "",
        "python eval/eval_harness.py",
        "# Results: eval/results.json",
    ])

    doc.add_heading("Example Conversations", level=2)
    doc.add_heading("Example 1: Factual query with citations", level=3)
    add_para(doc, 'User: "What organization maintains the Unicode Standard?"')
    add_para(doc, 'Agent: "The Unicode Standard is maintained by the Unicode Consortium [1]. The consortium coordinates standardization work for text encoding across software systems [1]."')
    doc.add_heading("Example 2: Impossible query with insufficient-evidence refusal", level=3)
    add_para(doc, 'User: "What is the exact GDP of the fictional underwater city of Atlantis in 2024?"')
    add_para(doc, 'Agent: "I cannot find sufficient evidence to answer this. Atlantis is fictional, and the retrieved sources do not provide a reliable real-world GDP figure. A refined search could look for fictional portrayals of Atlantis economics or worldbuilding estimates."')

    doc.save(DOCX_PATH)


def main():
    MD_PATH.write_text(MARKDOWN, encoding="utf-8")
    build_docx()
    print(DOCX_PATH)
    print(MD_PATH)


if __name__ == "__main__":
    main()
